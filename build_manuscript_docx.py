from __future__ import annotations

import re
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = Path(os.environ.get("MANUSCRIPT_SOURCE", ROOT / "MANUSCRIPT.md"))
OUTPUT = Path(os.environ.get("MANUSCRIPT_OUTPUT", ROOT / "Manuscript.docx"))
RESULTS = ROOT / "results"
DOCX_FIGURES = Path(os.environ.get("MANUSCRIPT_FIGURES", ROOT / "_docx_figures"))
TEMPLATE = Path(os.environ.get("MANUSCRIPT_TEMPLATE", ROOT / "_template_reference" / "Manuscript_format_reference.docx"))

FIGURES = {
    1: DOCX_FIGURES / "pca_explained_variance.png",
    2: DOCX_FIGURES / "patient_prediction_summary.png",
    3: DOCX_FIGURES / "loss_curves_MLP.png",
    4: DOCX_FIGURES / "loss_curves_CNN.png",
    5: DOCX_FIGURES / "loss_curves_BiLSTM.png",
    6: DOCX_FIGURES / "loss_curves_Transformer.png",
    7: DOCX_FIGURES / "loss_curves_XGBoost.png",
    8: DOCX_FIGURES / "shap_summary_plot.png",
}

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)


def set_font(run, name="Calibri", size=11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Article Title" not in styles:
        style = styles.add_style("Article Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Article Title"]
    style.font.name = "Calibri"
    style.font.size = Pt(18)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(14)
    style.paragraph_format.keep_with_next = True

    if "Caption Text" not in styles:
        cap = styles.add_style("Caption Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Caption Text"]
    cap.font.name = "Calibri"
    cap.font.size = Pt(10)
    cap.font.italic = False
    cap._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    cap._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.line_spacing = 1.0

    if "Reference" not in styles:
        ref = styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = styles["Reference"]
    ref.font.name = "Calibri"
    ref.font.size = Pt(10)
    ref._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    ref._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    ref.paragraph_format.left_indent = Inches(0.25)
    ref.paragraph_format.first_line_indent = Inches(-0.25)
    ref.paragraph_format.space_after = Pt(5)
    ref.paragraph_format.line_spacing = 1.0


def add_field(paragraph, instruction: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    for node in (begin, instr, separate, text, end):
        run.append(node)


def configure_page(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("Manuscript #93768"), size=8.5, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    set_font(footer.add_run("Page "), size=8.5, color=GRAY)
    add_field(footer, "PAGE")


INLINE = re.compile(r"(\*\*.*?\*\*|`.*?`)")


def add_inline(paragraph, text: str):
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Consolas", size=9.5)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row.height = None
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def widths_for(headers):
    n = len(headers)
    if n == 7:
        return [1050, 1350, 850, 900, 850, 2180, 2180]
    if n == 6:
        if headers[0].strip().lower() == "analysis":
            return [1900, 950, 1850, 1400, 1600, 1660]
        return [1250, 1950, 1550, 1550, 1550, 1510]
    if n == 5:
        return [1700, 1915, 1915, 1915, 1915]
    if n == 4:
        return [1200, 1800, 2400, 3960]
    return [9360 // n] * n


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths_for(rows[0]))
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            cell = table.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            if c > 0 and (r == 0 or re.fullmatch(r"[-+0-9.,()%]+", value.strip())):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, value)
            for run in p.runs:
                set_font(run, size=9.5, bold=(r == 0))
            if r == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F2F4F7")
                cell._tc.get_or_add_tcPr().append(shd)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_figure(doc, number: int, caption: str):
    path = FIGURES[number]
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(6.15))
    doc_pr = shape._inline.docPr
    doc_pr.set("title", f"Figure {number}")
    doc_pr.set("descr", caption)
    cp = doc.add_paragraph(style="Caption Text")
    add_inline(cp, caption)


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document(TEMPLATE) if TEMPLATE.exists() else Document()
    if TEMPLATE.exists():
        body = doc._element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)
    configure_styles(doc)
    if not TEMPLATE.exists():
        configure_page(doc)

    in_references = False
    i = 0
    title_seen = False
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                values = [v.strip() for v in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", value) for value in values):
                    rows.append(values)
                i += 1
            add_table(doc, rows)
            continue

        figure_match = re.match(r"Figure (\d+)\.\s*(.*)", line)
        if figure_match:
            number = int(figure_match.group(1))
            add_figure(doc, number, line)
            i += 1
            continue

        if line.startswith("# "):
            text = line[2:].strip()
            if text == "Original Paper":
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(5)
                set_font(p.add_run(text), size=10, bold=True, color=GRAY)
            elif not title_seen:
                p = doc.add_paragraph(style="Article Title")
                add_inline(p, text)
                title_seen = True
            else:
                doc.add_paragraph(text, style="Heading 1")
                in_references = text == "References"
            i += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 2")
            i += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 3")
            i += 1
            continue

        if line.startswith("Table "):
            p = doc.add_paragraph(style="Caption Text")
            p.paragraph_format.keep_with_next = True
            add_inline(p, line)
            i += 1
            continue

        if in_references and re.match(r"\d+\.\s", line):
            p = doc.add_paragraph(style="Reference")
            add_inline(p, line)
            i += 1
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            candidate = lines[i].strip()
            if not candidate or candidate.startswith(("#", "|", "Table ", "Figure ")):
                break
            if in_references and re.match(r"\d+\.\s", candidate):
                break
            paragraph_lines.append(candidate)
            i += 1
        p = doc.add_paragraph()
        text = " ".join(paragraph_lines).replace("  ", " ")
        add_inline(p, text)

    doc.core_properties.title = "Baseline Peripheral Blood Single-Cell Multimodal Profiling for Chemoimmunotherapy Response"
    doc.core_properties.subject = "JMIR Bioinformatics and Biotechnology manuscript #93768"
    doc.core_properties.keywords = "breast cancer; immunotherapy; single-cell; TCR; machine learning"
    doc.core_properties.author = ""
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
